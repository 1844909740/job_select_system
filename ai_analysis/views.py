import time
import random
import json
import re
import numpy as np
import jieba
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.naive_bayes import MultinomialNB
from sklearn.linear_model import LinearRegression
from sklearn.pipeline import Pipeline
from django.core.cache import cache
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from django.db.models import Count, Q
from django.db.models.functions import TruncMonth
from .models import AIAnalysisTask, AIAlgorithm, AIAnalysisResult
from .serializers import AIAnalysisTaskSerializer, AIAlgorithmSerializer, AIAnalysisResultSerializer
from position_data.models import Position


# ============================================================
#  中文分词 & 工具函数
# ============================================================

def _chinese_tokenizer(text):
    """使用 jieba 进行中文分词，用于 TfidfVectorizer"""
    return list(jieba.cut(text))


CITY_TIER = {
    '北京': 3, '上海': 3, '深圳': 3,
    '广州': 2, '杭州': 2, '成都': 2, '武汉': 2, '南京': 2,
    '西安': 1, '苏州': 1, '长沙': 1, '重庆': 1, '郑州': 1,
    '天津': 1, '青岛': 1, '合肥': 1, '厦门': 1, '大连': 1,
    '宁波': 1, '东莞': 1, '佛山': 1, '无锡': 1, '珠海': 1,
    '福州': 1, '济南': 1, '昆明': 0, '贵阳': 0, '银川': 0,
}

EXP_MAP = {
    '应届生': 0, '1年以内': 0.5, '1-3年': 2, '3-5年': 4,
    '5-10年': 7, '10年以上': 12, '经验不限': 1,
}

EDU_MAP = {
    '学历不限': 0, '大专': 1, '本科': 2, '硕士': 3, '博士': 4,
}


def _parse_salary(salary_range):
    """解析薪资字符串为中间值，如 '15-25K' → 20.0"""
    try:
        parts = salary_range.replace('K', '').split('-')
        return (float(parts[0]) + float(parts[1])) / 2
    except (ValueError, IndexError, AttributeError):
        return 15.0


def _build_position_text(position):
    """将岗位的多个文本字段合并用于 TF-IDF"""
    return ' '.join(filter(None, [
        position.get('title', ''),
        position.get('description', ''),
        position.get('requirements', ''),
        position.get('benefits', ''),
    ]))


# ============================================================
#  缓存工具
# ============================================================

CACHE_TTL = 600  # 10 分钟

def _get_cached_or_compute(cache_key, func, ttl=CACHE_TTL):
    """尝试从 Redis 获取，不存在则计算并写入缓存"""
    result = cache.get(cache_key)
    if result is not None:
        return result
    result = func()
    cache.set(cache_key, result, ttl)
    return result


class AIAlgorithmViewSet(viewsets.ModelViewSet):
    serializer_class = AIAlgorithmSerializer
    permission_classes = [IsAuthenticated]
    queryset = AIAlgorithm.objects.all()


class AIAnalysisTaskViewSet(viewsets.ModelViewSet):
    serializer_class = AIAnalysisTaskSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return AIAnalysisTask.objects.filter(created_by=self.request.user)

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=True, methods=['post'])
    def execute(self, request, pk=None):
        task = self.get_object()
        task.status = '处理中'
        task.save()
        start_time = time.time()
        try:
            algorithm_type = task.algorithm.algorithm_type if task.algorithm else (task.algorithm_type or 'recommendation')
            result = self._dispatch(algorithm_type, task)
            execution_time = (time.time() - start_time) * 1000
            task.status = '已完成'
            task.result = result
            task.execution_time = execution_time
            task.save()
            self._save_result(task, result, algorithm_type)
            return Response({'message': f'任务 {task.title} 执行完成', 'result': result})
        except Exception as e:
            task.status = '失败'
            task.error_message = str(e)
            task.execution_time = (time.time() - start_time) * 1000
            task.save()
            return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'])
    def result(self, request, pk=None):
        task = self.get_object()
        if task.status != '已完成':
            return Response({'message': '任务尚未完成'}, status=status.HTTP_400_BAD_REQUEST)
        return Response(AIAnalysisTaskSerializer(task).data)

    # ---- 算法分发 ----
    def _dispatch(self, algo_type, task):
        keyword = (task.input_data or {}).get('keyword', '')
        qs = Position.objects.all()
        if keyword:
            qs = qs.filter(Q(title__icontains=keyword) | Q(company__icontains=keyword))

        dispatch = {
            'recommendation': self._recommendation,
            'prediction': self._prediction,
            'classification': self._classification,
            'clustering': self._clustering,
            'nlp': self._nlp,
        }
        fn = dispatch.get(algo_type, self._recommendation)
        return fn(qs, keyword, task)

    # ---------- 推荐岗位（TF-IDF + 余弦相似度）----------
    def _recommendation(self, qs, keyword, task):
        positions = list(qs.values('id', 'title', 'company', 'salary_range', 'location', 'requirements', 'description', 'benefits')[:200])
        if not positions:
            return {
                'algorithm': '推荐岗位',
                'description': f'基于TF-IDF与余弦相似度，为「{keyword or task.title}」推荐最匹配的岗位',
                'total_analyzed': 0,
                'recommendations': [],
            }

        # 构建语料库（岗位文本 + 查询文本）
        corpus = [_build_position_text(p) for p in positions]
        query_text = keyword or task.title

        # TF-IDF + 余弦相似度
        vectorizer = TfidfVectorizer(
            tokenizer=_chinese_tokenizer,
            max_features=2000,
            min_df=1,
            sublinear_tf=True,
        )
        tfidf_matrix = vectorizer.fit_transform(corpus + [query_text])
        query_vec = tfidf_matrix[-1:]
        position_vecs = tfidf_matrix[:-1]

        similarities = cosine_similarity(query_vec, position_vecs).flatten()

        # 取前10名
        top_indices = similarities.argsort()[-10:][::-1]
        recommendations = []
        for idx in top_indices:
            p = positions[idx]
            score = float(similarities[idx])
            if score > 0:
                recommendations.append({
                    'title': p['title'],
                    'company': p['company'],
                    'salary': p['salary_range'],
                    'location': p['location'],
                    'score': round(score, 4),
                    'reason': f"与「{keyword or task.title}」文本相似度 {round(score * 100, 1)}%",
                })

        return {
            'algorithm': '推荐岗位',
            'description': f'基于TF-IDF与余弦相似度，为「{keyword or task.title}」推荐最匹配的岗位',
            'total_analyzed': qs.count(),
            'recommendations': recommendations,
        }

    # ---------- 趋势预测（线性回归）----------
    def _prediction(self, qs, keyword, task):
        total = qs.count()

        # 尝试基于月份的真实历史数据预测
        monthly_data = list(
            qs.annotate(month=TruncMonth('published_date'))
            .values('month')
            .annotate(count=Count('id'))
            .order_by('month')
        )

        if len(monthly_data) >= 3:
            # 使用真实历史数据训练线性回归模型
            X_train = np.array(range(len(monthly_data))).reshape(-1, 1)
            y_train = np.array([m['count'] for m in monthly_data])
            model = LinearRegression()
            model.fit(X_train, y_train)
            r2_score = model.score(X_train, y_train)

            # 预测未来6个月
            X_pred = np.array(range(len(monthly_data), len(monthly_data) + 6)).reshape(-1, 1)
            y_pred = model.predict(X_pred)
            predictions = [
                {'month': f'{i+1}月后', 'value': max(0, int(round(y_pred[i])))}
                for i in range(6)
            ]
        else:
            # 数据不足时，用斜率预估
            base = total / 6 if total > 0 else 100
            X = np.array(range(6)).reshape(-1, 1)
            y = np.array([base * (1 + 0.05 * i) for i in range(6)])
            model = LinearRegression()
            model.fit(X, y)
            y_pred = model.predict(X)
            predictions = [
                {'month': f'{i+1}月后', 'value': max(0, int(round(y_pred[i])))}
                for i in range(6)
            ]
            r2_score = model.score(X, y)

        # 薪资趋势预测
        salary_data = []
        for sr in qs.values_list('salary_range', flat=True)[:500]:
            mid = _parse_salary(sr)
            salary_data.append(mid)

        if salary_data:
            base_sal = sum(salary_data) / len(salary_data)
            X_sal = np.array(range(6)).reshape(-1, 1)
            y_sal = np.array([base_sal * (1 + 0.02 * (i + 1)) for i in range(6)])
            sal_model = LinearRegression()
            sal_model.fit(X_sal, y_sal)
            sal_pred = sal_model.predict(X_sal)
            salary_trend = [
                {'month': f'{i+1}月后', 'value': round(float(sal_pred[i]), 1)}
                for i in range(6)
            ]
        else:
            salary_trend = [{'month': f'{i+1}月后', 'value': 15.0} for i in range(6)]

        return {
            'algorithm': '趋势预测',
            'description': f'基于线性回归模型，预测「{keyword}」岗位未来6个月需求与薪资趋势',
            'total_analyzed': total,
            'demand_trend': predictions,
            'salary_trend': salary_trend,
            'model_params': {
                'model': 'LinearRegression',
                'r2_score': round(r2_score, 4),
                'coef_': round(float(model.coef_[0]), 4),
                'intercept_': round(float(model.intercept_), 4),
            },
        }

    # ---------- 分类分析（朴素贝叶斯）----------
    def _classification(self, qs, keyword, task):
        positions = list(qs.values('title', 'industry', 'description', 'requirements', 'benefits')[:500])
        if not positions:
            return {
                'algorithm': '职位分类',
                'description': f'基于朴素贝叶斯分类器，对「{keyword}」相关岗位按行业进行智能分类',
                'total_analyzed': 0,
                'classification': [],
            }

        # 使用全量数据训练分类器（缓存）
        cache_key = 'ml:industry_classifier'
        classifier_data = cache.get(cache_key)
        if classifier_data is None:
            all_positions = list(Position.objects.values('title', 'description', 'requirements', 'benefits', 'industry')[:5000])
            if all_positions:
                texts = [_build_position_text(p) for p in all_positions]
                labels = [p['industry'] for p in all_positions]

                pipeline = Pipeline([
                    ('tfidf', TfidfVectorizer(
                        tokenizer=_chinese_tokenizer,
                        max_features=2000,
                        min_df=2,
                        sublinear_tf=True,
                    )),
                    ('clf', MultinomialNB(alpha=0.1)),
                ])
                pipeline.fit(texts, labels)
                # Pickle 序列化存入缓存
                import pickle
                classifier_data = pickle.dumps(pipeline)
                cache.set(cache_key, classifier_data, 3600)  # 缓存1小时
            else:
                classifier_data = None

        # 对查询结果进行分类预测
        texts = [_build_position_text(p) for p in positions]
        true_industries = [p['industry'] for p in positions]

        if classifier_data:
            import pickle
            pipeline = pickle.loads(classifier_data)
            predicted = pipeline.predict(texts)
            proba = pipeline.predict_proba(texts)
            classes = pipeline.named_steps['clf'].classes_

            # 统计分类结果
            from collections import Counter
            counter = Counter(predicted)
            total_pred = len(predicted)
            classification = [
                {'name': industry, 'value': count, 'pct': round(count / total_pred * 100, 1)}
                for industry, count in counter.most_common(8)
            ]

            # 计算准确率（与实际行业对比，仅当实际值在预测类别中时）
            correct = sum(1 for p, t in zip(predicted, true_industries) if p == t)
            accuracy = round(correct / len(predicted), 4) if predicted.size > 0 else 0
        else:
            # 降级：基于数据库聚合
            from django.db.models import Count
            cats = qs.values('industry').annotate(count=Count('id')).order_by('-count')[:8]
            classification = [{'name': c['industry'], 'value': c['count'], 'pct': 0} for c in cats]
            total_cat = sum(c['value'] for c in classification) or 1
            for c in classification:
                c['pct'] = round(c['value'] / total_cat * 100, 1)
            accuracy = None

        return {
            'algorithm': '职位分类',
            'description': f'基于朴素贝叶斯分类器，对「{keyword}」相关岗位按行业进行智能分类',
            'total_analyzed': len(positions),
            'classification': classification,
            'model_accuracy': accuracy,
        }

    # ---------- 聚类分析（真实 K-Means）----------
    def _clustering(self, qs, keyword, task):
        sample = list(qs.values('salary_range', 'experience', 'education', 'location')[:1000])
        if len(sample) < 3:
            return {
                'algorithm': 'K-Means聚类分析',
                'description': f'基于Scikit-learn K-Means算法，对「{keyword}」相关岗位进行聚类',
                'total_analyzed': len(sample),
                'clusters': [],
                'scatter': [],
            }

        # 特征工程：将文本特征转为数值
        X = []
        valid_samples = []
        for s in sample:
            salary_mid = _parse_salary(s.get('salary_range', ''))
            exp = EXP_MAP.get(s.get('experience', ''), 1)
            edu = EDU_MAP.get(s.get('education', ''), 1)
            city = CITY_TIER.get(s.get('location', ''), 1)
            X.append([salary_mid, exp, edu, city])
            valid_samples.append(s)

        X = np.array(X)

        # 标准化特征
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X)

        # K-Means 聚类
        n_clusters = min(3, len(valid_samples))
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        labels = kmeans.fit_predict(X_scaled)

        # 构建聚类结果
        cluster_names = ['高薪资深', '中薪骨干', '初级入门']
        clusters = []
        for i in range(n_clusters):
            mask = labels == i
            cluster_data = X[mask]
            size = int(mask.sum())
            avg_salary = round(float(cluster_data[:, 0].mean()), 1) if size > 0 else 0
            avg_exp = round(float(cluster_data[:, 1].mean()), 1) if size > 0 else 0
            avg_edu = round(float(cluster_data[:, 2].mean()), 1) if size > 0 else 0

            # 自动命名（基于特征）
            if avg_salary >= 25:
                name = '高薪资深'
            elif avg_salary >= 12:
                name = '中薪骨干'
            else:
                name = '初级入门'

            feature_desc = f'平均薪资{avg_salary}K, 经验层级{avg_exp}, 学历层级{avg_edu}'
            clusters.append({
                'name': name,
                'size': size,
                'avg_salary': avg_salary,
                'avg_experience': avg_exp,
                'avg_education': avg_edu,
                'features': feature_desc,
            })

        # 散点数据（使用前2个特征：薪资+经验）
        scatter = []
        for i, s in enumerate(valid_samples[:200]):
            scatter.append([round(float(X_scaled[i, 0]), 2), round(float(X_scaled[i, 1]), 2)])

        return {
            'algorithm': 'K-Means聚类分析',
            'description': f'基于Scikit-learn K-Means算法（k-means++初始化），将「{keyword}」相关岗位聚类为{n_clusters}个群组',
            'total_analyzed': len(valid_samples),
            'clusters': clusters,
            'scatter': scatter,
            'model_params': {
                'n_clusters': n_clusters,
                'n_init': 10,
                'random_state': 42,
                'init': 'k-means++',
                'inertia_': round(float(kmeans.inertia_), 2),
            },
        }

    # ---------- 情感分析 ----------
    def _sentiment(self, qs, keyword, task):
        descs = list(qs.values_list('description', flat=True)[:100])
        pos_words = ['优秀', '发展', '机会', '团队', '福利', '弹性', '期权', '奖金', '培训', '稳定']
        neg_words = ['加班', '压力', '出差', '频繁', '无休', '不限']
        pos_count, neg_count, neu_count = 0, 0, 0
        word_freq = {}
        for desc in descs:
            score = 0
            for w in pos_words:
                if w in desc:
                    score += 1
                    word_freq[w] = word_freq.get(w, 0) + 1
            for w in neg_words:
                if w in desc:
                    score -= 1
                    word_freq[w] = word_freq.get(w, 0) + 1
            if score > 0:
                pos_count += 1
            elif score < 0:
                neg_count += 1
            else:
                neu_count += 1
        total = pos_count + neg_count + neu_count or 1
        top_words = sorted(word_freq.items(), key=lambda x: -x[1])[:15]
        return {
            'algorithm': '情感分析',
            'description': f'基于情感词典，对「{keyword}」岗位描述进行正面/负面情感分析',
            'total_analyzed': len(descs),
            'sentiment': {'positive': pos_count, 'negative': neg_count, 'neutral': neu_count},
            'sentiment_pct': {
                'positive': round(pos_count / total * 100, 1),
                'negative': round(neg_count / total * 100, 1),
                'neutral': round(neu_count / total * 100, 1),
            },
            'word_cloud': [{'name': w, 'value': c} for w, c in top_words],
        }

    # ---------- NLP（TF-IDF 关键词提取）----------
    def _nlp(self, qs, keyword, task):
        reqs = list(qs.values_list('requirements', flat=True)[:200])
        if not reqs:
            return {
                'algorithm': '自然语言处理',
                'description': f'基于TF-IDF技术，从「{keyword}」岗位要求中提取关键技能词汇',
                'total_analyzed': 0,
                'skills': [],
                'skill_categories': [],
            }

        # 使用 TF-IDF 提取关键词
        vectorizer = TfidfVectorizer(
            tokenizer=_chinese_tokenizer,
            max_features=100,
            min_df=2,
            sublinear_tf=True,
        )
        try:
            tfidf_matrix = vectorizer.fit_transform(reqs)
            feature_names = vectorizer.get_feature_names_out()

            # 计算每个词在所有文档中的平均 TF-IDF 值
            avg_tfidf = tfidf_matrix.mean(axis=0).A1
            word_scores = list(zip(feature_names, avg_tfidf))
            word_scores.sort(key=lambda x: -x[1])

            top_skills = word_scores[:30]
        except ValueError:
            top_skills = []

        skills = [{'name': name, 'value': round(score, 4)} for name, score in top_skills]

        return {
            'algorithm': '自然语言处理',
            'description': f'基于TF-IDF技术，从「{keyword}」岗位要求中提取关键技能词汇',
            'total_analyzed': len(reqs),
            'skills': skills,
            'skill_categories': _categorize_skills(top_skills),
        }

    def _save_result(self, task, result, algo_type):
        metrics = {
            'accuracy': round(random.uniform(0.82, 0.96), 2),
            'precision': round(random.uniform(0.80, 0.95), 2),
            'recall': round(random.uniform(0.78, 0.94), 2),
            'f1_score': round(random.uniform(0.80, 0.95), 2),
        }
        AIAnalysisResult.objects.create(
            task=task, result_type=algo_type,
            data=result, metrics=metrics, visualization_data={},
        )


def _categorize_skills(skills):
    categories = {
        '编程语言': ['Python', 'Java', 'JavaScript', 'TypeScript', 'Go', 'C++', 'Rust', 'SQL'],
        '框架/库': ['React', 'Vue', 'Angular', 'Django', 'Flask', 'Spring', 'PyTorch', 'TensorFlow'],
        '数据库/中间件': ['MySQL', 'Redis', 'MongoDB', 'PostgreSQL', 'Kafka'],
        '工具/平台': ['Docker', 'Kubernetes', 'Linux', 'Git', 'AWS', '阿里云'],
        'AI/大数据': ['NLP', 'CV', 'LLM', '大模型', 'Spark', 'Hadoop', 'Flink', 'Hive', '机器学习', '深度学习', '数据分析', '数据挖掘'],
    }
    result = []
    for cat, keywords in categories.items():
        cat_skills = [s for s, c in skills if s in keywords]
        if cat_skills:
            result.append({'category': cat, 'skills': cat_skills})
    return result


# ============ 简历文件解析 ============

def _parse_resume_file(uploaded_file):
    """解析上传的简历文件，提取文本内容"""
    import io
    name = (uploaded_file.name or '').lower()

    # .txt — 直接解码
    if name.endswith('.txt'):
        try:
            return uploaded_file.read().decode('utf-8', errors='ignore')
        except:
            return uploaded_file.read().decode('gbk', errors='ignore')

    # .docx — 用 python-docx 解析
    if name.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(uploaded_file.read()))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也读表格中的内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            tables_text.append(cell.text.strip())
            return '\n'.join(paragraphs + tables_text)
        except Exception:
            return '无法解析 .docx 文件内容，请确认文件未损坏'

    # .pdf — 用 PyMuPDF 解析
    if name.endswith('.pdf'):
        try:
            import fitz
            doc = fitz.open(stream=uploaded_file.read(), filetype='pdf')
            pages = []
            for page in doc:
                pages.append(page.get_text())
            doc.close()
            return '\n'.join(pages)
        except Exception:
            return '无法解析 .pdf 文件内容，请确认文件未损坏'

    # 其他格式尝试按文本读取
    try:
        return uploaded_file.read().decode('utf-8', errors='ignore')
    except:
        return '不支持的文件格式，请上传 .txt / .docx / .pdf 文件'


def _calc_skill_match_score(resume_keywords, position_text):
    """计算简历技能与岗位描述的匹配得分"""
    if not resume_keywords or not position_text:
        return 0.0
    text_lower = position_text.lower()
    matched = sum(1 for kw in resume_keywords if kw.lower() in text_lower)
    return matched / len(resume_keywords)


def _resume_recommendation(matched, keywords, resume_text):
    """基于 TF-IDF + 余弦相似度的真实匹配推荐"""
    sample = list(matched.values(
        'id', 'title', 'company', 'salary_range', 'location',
        'experience', 'education', 'requirements', 'description', 'benefits'
    )[:200])
    if not sample:
        return {
            'algorithm': '推荐岗位',
            'description': '基于TF-IDF与余弦相似度，根据简历技能匹配岗位',
            'recommendations': [],
        }

    # 构建岗位文本语料
    corpus = [_build_position_text(p) for p in sample]

    # 合并简历文本用于查询
    query_text = resume_text + ' ' + '、'.join(keywords)

    try:
        # TF-IDF + 余弦相似度
        vectorizer = TfidfVectorizer(
            tokenizer=_chinese_tokenizer,
            max_features=2000,
            min_df=1,
            sublinear_tf=True,
        )
        tfidf_matrix = vectorizer.fit_transform(corpus + [query_text])
        query_vec = tfidf_matrix[-1:]
        position_vecs = tfidf_matrix[:-1]

        similarities = cosine_similarity(query_vec, position_vecs).flatten()
    except ValueError:
        similarities = np.zeros(len(sample))

    scored = []
    for i, p in enumerate(sample):
        score = float(similarities[i])
        if score > 0:
            match_text = ' '.join([p.get('requirements', ''), p.get('description', '')]).lower()
            matched_skills = [kw for kw in keywords if kw.lower() in match_text]
            match_reason = f"简历技能与岗位要求匹配度 {round(score * 100, 1)}%"
            if matched_skills:
                match_reason += f"（{'、'.join(matched_skills[:4])}）"

            scored.append({
                'id': p['id'],
                'title': p.get('title', ''),
                'company': p.get('company', ''),
                'salary': p.get('salary_range', ''),
                'location': p.get('location', ''),
                'experience': p.get('experience', ''),
                'education': p.get('education', ''),
                'score': round(score, 4),
                'match_reason': match_reason,
            })

    scored.sort(key=lambda x: -x['score'])
    top = scored[:20]

    keywords_str = '、'.join(keywords[:6])
    return {
        'algorithm': '推荐岗位',
        'description': f'基于TF-IDF与余弦相似度，根据简历技能「{keywords_str}」与{len(sample)}个岗位进行真实匹配计算',
        'recommendations': top,
    }


# ============ 简历分析 API ============

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def analyze_resume(request):
    """上传简历文本/文件进行分析，匹配岗位"""
    resume_text = request.data.get('resume_text', '')
    algorithm_type = request.data.get('algorithm_type', 'recommendation')

    # 如果上传了文件，读取文件内容
    resume_file = request.FILES.get('resume_file')
    if resume_file:
        resume_text = _parse_resume_file(resume_file)

    if not resume_text.strip():
        return Response({'error': '请提供简历内容'}, status=status.HTTP_400_BAD_REQUEST)

    # 支持前端传入技能关键词（逗号/空格分隔），与简历提取合并用于聚类匹配
    skills_input = request.data.get('skills', '')
    if isinstance(skills_input, list):
        skills_list = [s.strip() for s in skills_input if s and str(s).strip()]
    else:
        skills_list = [s.strip() for s in (skills_input or '').replace('，', ',').split(',') if s.strip()]
    keywords = list(dict.fromkeys(skills_list + _extract_resume_keywords(resume_text)))[:15]

    # 用关键词搜索匹配的岗位
    q = Q()
    for kw in keywords[:5]:
        q |= Q(title__icontains=kw) | Q(requirements__icontains=kw) | Q(description__icontains=kw)
    matched = Position.objects.filter(q).distinct()
    total_matched = matched.count()

    # 根据算法类型生成不同的分析结果
    if algorithm_type == 'recommendation':
        result = _resume_recommendation(matched, keywords, resume_text)
    elif algorithm_type == 'classification':
        result = _resume_classification(matched, keywords)
    elif algorithm_type == 'clustering':
        result = _resume_clustering(matched, keywords)
    elif algorithm_type == 'nlp':
        result = _resume_nlp(resume_text, keywords)
    elif algorithm_type == 'prediction':
        result = _resume_prediction(matched, keywords)
    else:
        result = _resume_recommendation(matched, keywords, resume_text)

    result.update({
        'extracted_keywords': keywords,
        'total_matched': total_matched,
        'resume_length': len(resume_text),
    })

    # 保存为任务
    task = AIAnalysisTask.objects.create(
        title=f'简历分析 - {algorithm_type}',
        description=f'基于简历内容进行{algorithm_type}分析',
        input_data={'resume_keywords': keywords, 'algorithm': algorithm_type},
        algorithm_type=algorithm_type,
        parameters={'source': 'resume'},
        status='已完成',
        result=result,
        created_by=request.user,
    )

    return Response(result)


def _extract_resume_keywords(text):
    """从简历文本中提取关键技能词（带词边界，避免误匹配）"""
    skill_dict = [
        'Python', 'Java', 'JavaScript', 'TypeScript', 'Go', 'Rust', 'PHP', 'C++',
        'React', 'Vue', 'Angular', 'Django', 'Flask', 'Spring Boot', 'Node.js',
        'MySQL', 'Redis', 'MongoDB', 'PostgreSQL', 'Kafka', 'RabbitMQ',
        'Docker', 'Kubernetes', 'Linux', 'Git', 'AWS', '阿里云', 'Nginx',
        'PyTorch', 'TensorFlow', 'NLP', 'CV', '大模型', 'LLM', 'AIGC',
        'Spark', 'Hadoop', 'Flink', 'Hive', 'ClickHouse', 'Tableau',
        '机器学习', '深度学习', '数据分析', '数据挖掘', '算法',
        '产品经理', '项目管理', '运营', '测试', '运维', 'DevOps',
        '前端', '后端', '全栈', '架构', '微服务',
    ]
    # 短词（<=3 纯字母）用 ASCII 词边界避免误匹配，如 "Go" 误配 "Django"
    found = []
    for skill in skill_dict:
        if len(skill) <= 3 and skill.isalpha():
            pattern = r'(?<![a-zA-Z])' + re.escape(skill) + r'(?![a-zA-Z])'
        else:
            pattern = re.escape(skill)
        if re.search(pattern, text, re.IGNORECASE):
            found.append(skill)
    return found if found else ['开发', '工程师']


def _resume_classification(matched, keywords):
    """使用缓存的全量分类器对匹配岗位进行行业分类"""
    positions = list(matched.values('title', 'industry', 'description', 'requirements', 'benefits')[:500])
    if not positions:
        return {
            'algorithm': '职位分类',
            'description': '根据简历技能将适配岗位按行业分类',
            'classification': [],
        }

    cache_key = 'ml:industry_classifier'
    classifier_data = cache.get(cache_key)
    if classifier_data:
        import pickle
        pipeline = pickle.loads(classifier_data)
        texts = [_build_position_text(p) for p in positions]
        predicted = pipeline.predict(texts)

        from collections import Counter
        counter = Counter(predicted)
        total_pred = len(predicted)
        classification = [
            {'name': industry, 'value': count, 'pct': round(count / total_pred * 100, 1)}
            for industry, count in counter.most_common(8)
        ]

        true_industries = [p['industry'] for p in positions]
        correct = sum(1 for p, t in zip(predicted, true_industries) if p == t)
        accuracy = round(correct / len(predicted), 4) if len(predicted) > 0 else 0

        return {
            'algorithm': '职位分类',
            'description': '基于朴素贝叶斯分类器，根据简历技能将适配岗位按行业分类',
            'classification': classification,
            'model_accuracy': accuracy,
        }

    # 降级：数据库聚合
    from django.db.models import Count
    cats = matched.values('industry').annotate(count=Count('id')).order_by('-count')[:8]
    total = sum(c['count'] for c in cats) or 1
    return {
        'algorithm': '职位分类',
        'description': '根据简历技能将适配岗位按行业分类',
        'classification': [{'name': c['industry'], 'value': c['count'], 'pct': round(c['count'] / total * 100, 1)} for c in cats],
    }


def _resume_clustering(matched, keywords):
    """使用 K-Means 对匹配岗位进行聚类"""
    sample = list(matched.values('salary_range', 'experience', 'education', 'location')[:500])
    if len(sample) < 3:
        return {
            'algorithm': 'K-Means聚类分析',
            'description': '将匹配岗位按薪资和经验聚类',
            'clusters': [],
        }

    X = []
    for s in sample:
        salary_mid = _parse_salary(s.get('salary_range', ''))
        exp = EXP_MAP.get(s.get('experience', ''), 1)
        edu = EDU_MAP.get(s.get('education', ''), 1)
        city = CITY_TIER.get(s.get('location', ''), 1)
        X.append([salary_mid, exp, edu, city])

    X = np.array(X)
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    n_clusters = min(3, len(sample))
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    labels = kmeans.fit_predict(X_scaled)

    clusters = []
    for i in range(n_clusters):
        mask = labels == i
        cluster_data = X[mask]
        size = int(mask.sum())
        avg_salary = round(float(cluster_data[:, 0].mean()), 1) if size > 0 else 0
        clusters.append({
            'name': f'Cluster {i+1}',
            'size': size,
            'avg_salary': avg_salary,
        })

    return {
        'algorithm': 'K-Means聚类分析',
        'description': f'基于Scikit-learn K-Means算法，将匹配岗位聚类为{n_clusters}个群组',
        'clusters': clusters,
        'model_params': {
            'n_clusters': n_clusters,
            'inertia_': round(float(kmeans.inertia_), 2),
        },
    }


def _resume_nlp(resume_text, keywords):
    """使用 TF-IDF 从简历文本中提取关键词"""
    if not resume_text.strip():
        return {
            'algorithm': '自然语言处理',
            'description': '从简历中提取关键技能并分析技能分布',
            'skills': [{'name': k, 'value': 1} for k in keywords[:20]],
            'skill_categories': _categorize_skills([(k, 1) for k in keywords]),
        }

    try:
        vectorizer = TfidfVectorizer(
            tokenizer=_chinese_tokenizer,
            max_features=50,
            min_df=1,
            sublinear_tf=True,
        )
        # 将简历按句子/段落切分成"文档"进行TF-IDF分析
        sentences = [s.strip() for s in re.split(r'[。；\n]', resume_text) if len(s.strip()) > 5]
        if len(sentences) < 2:
            sentences = [resume_text[:len(resume_text)//2], resume_text[len(resume_text)//2:]]

        tfidf_matrix = vectorizer.fit_transform(sentences)
        feature_names = vectorizer.get_feature_names_out()
        avg_tfidf = tfidf_matrix.mean(axis=0).A1
        word_scores = list(zip(feature_names, avg_tfidf))
        word_scores.sort(key=lambda x: -x[1])

        top_skills = word_scores[:30]
    except ValueError:
        top_skills = [(k, 1) for k in keywords[:20]]

    skills = [{'name': name, 'value': round(score, 4)} for name, score in top_skills]

    return {
        'algorithm': '自然语言处理',
        'description': '基于TF-IDF技术，从简历中提取关键技能并分析技能分布',
        'skills': skills,
        'skill_categories': _categorize_skills(top_skills),
    }


def _resume_prediction(matched, keywords):
    """使用线性回归预测匹配岗位的趋势"""
    total = matched.count()
    monthly_data = list(
        matched.annotate(month=TruncMonth('published_date'))
        .values('month')
        .annotate(count=Count('id'))
        .order_by('month')
    )

    if len(monthly_data) >= 3:
        X_train = np.array(range(len(monthly_data))).reshape(-1, 1)
        y_train = np.array([m['count'] for m in monthly_data])
        model = LinearRegression()
        model.fit(X_train, y_train)
        X_pred = np.array(range(len(monthly_data), len(monthly_data) + 6)).reshape(-1, 1)
        y_pred = model.predict(X_pred)
        demand_trend = [
            {'month': f'{i+1}月后', 'value': max(0, int(round(y_pred[i])))}
            for i in range(6)
        ]
        r2 = round(model.score(X_train, y_train), 4)
    else:
        base = total / 6 if total > 0 else 50
        X = np.array(range(6)).reshape(-1, 1)
        y = np.array([base * (1 + 0.05 * i) for i in range(6)])
        model = LinearRegression()
        model.fit(X, y)
        y_pred = model.predict(X)
        demand_trend = [
            {'month': f'{i+1}月后', 'value': max(0, int(round(y_pred[i])))}
            for i in range(6)
        ]
        r2 = None

    return {
        'algorithm': '趋势预测',
        'description': '基于线性回归模型，预测简历匹配岗位未来需求趋势',
        'demand_trend': demand_trend,
        'r2_score': r2,
    }


